from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path('/Users/rogerio.bittencourt/KeyNotesAI')
ASSET = ROOT / 'docs/manual/assets'
OUT = ROOT / 'output/manual/Manual-KeyNotesAI.docx'

INK = '191A18'; GOLD = 'C89A56'; GREEN = '35694F'; PALE = 'F5F1E8'; MUTED = '73756F'; WHITE = 'FFFFFF'; LINE = 'DDD8CE'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None: node = OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def font(run, size=None, bold=None, color=None, name='Aptos'):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),name); run._element.rPr.rFonts.set(qn('w:hAnsi'),name)
    if size: run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if color: run.font.color.rgb=RGBColor.from_string(color)
    return run

def add_text(doc, text='', bold=False, color=INK, size=9.5, after=4, align=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.12
    if align: p.alignment=align
    font(p.add_run(text),size,bold,color); return p

def h(doc, text, level=1):
    p=doc.add_paragraph(style=f'Heading {level}'); p.add_run(text); return p

def bullets(doc, items, numbered=False):
    for i,item in enumerate(items,1):
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.first_line_indent=Inches(-.18); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.1
        prefix=f'{i}. ' if numbered else '• '
        font(p.add_run(prefix),9.5,True,GOLD); font(p.add_run(item),9.5,False,INK)

def callout(doc, title, body, color=GREEN):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(6.5)
    c=t.cell(0,0); shade(c, PALE); set_cell_margins(c,130,170,130,170)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(3); font(p.add_run(title+'  '),10,True,color); font(p.add_run(body),9.4,False,INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def step(doc, n, title, body):
    t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(.55); t.columns[1].width=Inches(5.95)
    a,b=t.rows[0].cells; shade(a,GREEN); shade(b,PALE); set_cell_margins(a,110,90,110,90); set_cell_margins(b,110,150,110,150)
    a.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=a.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run(str(n)),15,True,WHITE)
    p=b.paragraphs[0]; p.paragraph_format.space_after=Pt(2); font(p.add_run(title),10.5,True,INK)
    p=b.add_paragraph(); p.paragraph_format.space_after=Pt(0); font(p.add_run(body),9.3,False,INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def picture(doc, filename, caption, width=6.45):
    path=ASSET/filename
    if path.exists():
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2); p.add_run().add_picture(str(path),width=Inches(width))
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(7); font(p.add_run(caption),8,False,MUTED)

def page(doc): doc.add_page_break()

doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.27); sec.page_height=Inches(11.69); sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.6); sec.left_margin=Inches(.8); sec.right_margin=Inches(.8); sec.header_distance=Inches(.3); sec.footer_distance=Inches(.3)
styles=doc.styles
normal=styles['Normal']; normal.font.name='Aptos'; normal.font.size=Pt(9.5); normal.font.color.rgb=RGBColor.from_string(INK); normal.paragraph_format.space_after=Pt(4); normal.paragraph_format.line_spacing=1.12
for name,size,color,before,after in [('Heading 1',20,INK,12,7),('Heading 2',14,GREEN,10,5),('Heading 3',11,GOLD,7,3)]:
    s=styles[name]; s.font.name='Aptos Display'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
cap=styles.add_style('Manual Caption',WD_STYLE_TYPE.PARAGRAPH); cap.font.name='Aptos'; cap.font.size=Pt(8); cap.font.color.rgb=RGBColor.from_string(MUTED)

# cabeçalho/rodapé
header=sec.header; p=header.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(p.add_run('KEYNOTESAI  •  MANUAL DO USUÁRIO'),7.5,True,MUTED)
footer=sec.footer; p=footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('INOVALAB · IFSC Câmpus Florianópolis-Continente'),7.5,False,MUTED)

# capa
for _ in range(3): doc.add_paragraph()
logo=ROOT/'public/keynotesai-logo.png'
if logo.exists():
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(logo),width=Inches(1.35))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(16); font(p.add_run('KeyNotes'),34,True,INK,'Aptos Display'); font(p.add_run('AI'),34,True,GOLD,'Aptos Display')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(20); font(p.add_run('MEETING INTELLIGENCE'),10,True,MUTED)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8); font(p.add_run('Manual completo do usuário'),25,True,INK,'Aptos Display')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('Do acesso à reunião, documentos, Google Drive e Trello'),12,False,GREEN)
for _ in range(5): doc.add_paragraph()
callout(doc,'Versão 1.2','Guia prático para usuários autorizados e administradores. Atualizado em 24/08/2026, com identificação de locutores e conversa semântica.')
add_text(doc,'Acesse: keynotes-ai.rogerio-bittencourt.chatgpt.site',bold=True,color=GREEN,size=10,align=WD_ALIGN_PARAGRAPH.CENTER)

page(doc); h(doc,'Comece por aqui',1)
add_text(doc,'O KeyNotesAI transforma uma reunião em registro pesquisável: gravação, transcrição, resumo, ata, decisões, ações e arquivos organizados. O fluxo recomendado abaixo cobre o uso completo.',size=10.5)
for n,title,body in [
 (1,'Entrar','Acesse com a conta ChatGPT cujo e-mail foi previamente autorizado.'),
 (2,'Criar a reunião','Informe um título e escolha o processamento híbrido (padrão), totalmente OpenAI ou OpenAI + locutores.'),
 (3,'Escolher a origem e gravar','Use Google Meet + microfone para reuniões on-line, somente microfone para encontros presenciais, ou importe um áudio existente.'),
 (4,'Completar o registro','Revise dados, participantes e, se desejar, adicione uma foto separada da ata.'),
 (5,'Transcrever e analisar','Revise a transcrição e gere resumo, ações, decisões e documentos.'),
 (6,'Ajustar responsabilidades','Defina responsável, prazo, prioridade e status das ações.'),
 (7,'Arquivar no Drive','Crie ou atualize a pasta única da reunião no Google Drive institucional.'),
 (8,'Sincronizar Trello','Após o Drive, o mesmo card é criado ou atualizado automaticamente; use o botão manual para alterações posteriores.')]: step(doc,n,title,body)
callout(doc,'Regra de ouro','Primeiro revise os dados e documentos; depois arquive/atualize no Google Drive. Assim, os links corretos seguem para o Trello.')

page(doc); h(doc,'1. Acesso e navegação',1)
h(doc,'Entrar',2); bullets(doc,['Abra o endereço do KeyNotesAI.','Clique em “Entrar com ChatGPT”.','Escolha a conta cujo e-mail está autorizado pelo administrador.','Se aparecer “acesso não autorizado”, peça ao administrador para cadastrar exatamente o mesmo e-mail usado no ChatGPT.'],numbered=True)
picture(doc,'01-acesso.png','Tela de acesso. A conta ChatGPT serve apenas para autenticação; o consumo da API é controlado pelo projeto.')
page(doc); h(doc,'Menu principal',2); add_text(doc,'Visão geral inicia reuniões e resume o trabalho; Reuniões mostra a gravação em andamento; Arquivos reúne gravações e documentos; Ações e Decisões consolidam resultados; Pergunte à IA conversa com uma reunião; Administração controla usuários; Integrações mostra Trello e Drive.')
picture(doc,'02-visao-geral.png','Visão geral e menu lateral com as integrações conectadas.')

page(doc); h(doc,'2. Criar, gravar ou importar uma reunião',1)
h(doc,'Nova gravação',2); bullets(doc,['Na Visão geral, informe um título claro — ele identificará histórico, documentos, pasta e card.','Escolha a origem do áudio: Google Meet + microfone (recomendado), somente Google Meet ou somente microfone.','Escolha o modo de processamento. O padrão é Híbrido.','Clique no botão para iniciar a gravação.','Autorize a captura solicitada pelo navegador.','Durante a reunião, acompanhe o tempo e registre participantes quando necessário.','Ao terminar, encerre a gravação e aguarde o arquivo ser preparado.'],numbered=True)
h(doc,'Importar áudio',2); bullets(doc,['Clique em “Importar áudio”.','Selecione o arquivo salvo no computador.','Informe ou confirme o título e o modo de processamento.','Prossiga para a biblioteca em Arquivos.'])
callout(doc,'Qual modo escolher?','Híbrido: padrão e menor custo. Totalmente OpenAI: envia também o áudio para transcrição. OpenAI + locutores: separa as vozes, tenta associá-las aos nomes da chamada inicial e permite revisão manual antes dos documentos.')
picture(doc,'03-reunioes.png','A área Reuniões exibe somente a reunião real em andamento.')

page(doc); h(doc,'Google Meet: captura completa',1)
add_text(doc,'Para registrar corretamente uma reunião on-line, abra o Google Meet e o KeyNotesAI no Chrome ou Edge. O modo recomendado combina o áudio remoto da aba com sua voz captada pelo microfone.',size=10.5)
h(doc,'Google Meet + microfone (recomendado)',2)
for n,title,body in [
 (1,'Prepare as duas abas','Entre na reunião pelo Google Meet e mantenha o KeyNotesAI aberto em outra aba do mesmo navegador.'),
 (2,'Escolha a origem','No KeyNotesAI, selecione “Google Meet + microfone (recomendado)”.'),
 (3,'Inicie a gravação','Clique em “Gravar reunião do Google Meet”. O navegador abrirá a janela de compartilhamento.'),
 (4,'Selecione a aba correta','Escolha especificamente a aba do Google Meet — não escolha uma janela ou a tela inteira.'),
 (5,'Compartilhe o áudio','Marque “Compartilhar áudio da guia” e confirme. Sem essa opção, as vozes remotas não serão gravadas.'),
 (6,'Autorize o microfone','Permita o acesso ao microfone para incluir sua própria voz com boa qualidade.'),
 (7,'Finalize','Encerre no KeyNotesAI. Se você interromper o compartilhamento da aba, a gravação será encerrada e salva automaticamente.')]: step(doc,n,title,body)
callout(doc,'Uso de fones','Recomendado no modo Google Meet + microfone: evita que o microfone recapture o som dos alto-falantes e reduz eco.')
h(doc,'Outros modos',2); bullets(doc,['Google Meet — somente áudio da aba: grava os participantes remotos, mas pode não registrar bem sua voz local.','Presencial — somente microfone: indicado para pessoas reunidas fisicamente no mesmo ambiente.'])
callout(doc,'Consentimento','Avise os participantes antes de iniciar e obtenha consentimento para gravação e processamento. O navegador sempre exige uma seleção manual da aba por segurança.')

h(doc,'Chamada inicial para identificar vozes',2)
bullets(doc,['Antes da discussão, peça que cada pessoa diga claramente: “Meu nome é [nome completo] e estou presente”.','Faça a chamada também no Google Meet; o nome exibido na tela não vem tecnicamente ligado ao áudio capturado.','Evite apresentações simultâneas e aguarde uma pessoa terminar antes da próxima.','Depois da transcrição, confira a associação entre Locutor A/B e os nomes registrados.'])

h(doc,'3. Dados, participantes e foto',1)
h(doc,'Completar os dados',2); bullets(doc,['Confirme título, data e hora.','Preencha setor/local, pauta e participantes, quando esses campos estiverem disponíveis.','Use nomes completos sempre que possível: isso melhora ata, ações e respostas do chat.'])
h(doc,'Foto da reunião',2); bullets(doc,['Use “Tirar foto” para abrir a câmera em dispositivos/navegadores compatíveis.','Use “Selecionar arquivo” para escolher uma imagem existente.','Confira a miniatura antes de salvar/arquivar.','A foto é armazenada como anexo separado no Drive e, posteriormente, referenciada no Trello. Ela não entra na ata.'])
picture(doc,'11-foto-reuniao.png','Registro fotográfico separado dos documentos oficiais.')
callout(doc,'Privacidade','Registre a imagem somente com o conhecimento e consentimento dos participantes.')

page(doc); h(doc,'4. Transcrição e análise semântica',1)
h(doc,'Processar a reunião',2); bullets(doc,['Abra Arquivos e selecione a reunião.','Ouça o áudio se precisar confirmar a qualidade.','Inicie a transcrição. No modo híbrido, mantenha a aba aberta durante a etapa local.','Revise o texto: corrija nomes, siglas, números, datas e termos técnicos.','Salve a transcrição revisada.','Execute a análise semântica para gerar os resultados estruturados.'],numbered=True)
picture(doc,'04-arquivos.png','Biblioteca: selecione uma gravação para transcrever, analisar e gerar documentos.')
h(doc,'Identificação de locutores',2); bullets(doc,['Selecione “OpenAI + locutores” antes de transcrever.','O sistema separa as falas e adiciona horário e identificação provisória de locutor.','A chamada inicial e a lista de presença são usadas para sugerir os nomes.','Na área “Conferir identificação dos locutores”, corrija ou complete cada nome.','Somente depois da conferência, execute a análise e gere os documentos.'])
callout(doc,'Funciona em qualquer origem','Google Meet, reunião presencial e áudio importado são aceitos. A precisão depende da clareza do áudio, da distância do microfone e de não haver falas simultâneas. Quando não houver evidência suficiente, o sistema mantém Locutor A/B em vez de inventar um nome.')
h(doc,'O que a análise produz',2); bullets(doc,['Resumo e temas principais.','Participantes identificados.','Ações propostas, com responsável/prazo/prioridade quando mencionados.','Decisões, pendências e bloqueios.','Ata, resumo executivo, matriz de ações e registro de decisões.'])

page(doc); h(doc,'5. Documentos e histórico',1)
h(doc,'Revisar e editar',2); bullets(doc,['Em Arquivos, selecione a reunião processada.','Abra cada documento gerado.','Edite o conteúdo necessário e salve.','Baixe uma cópia local quando precisar compartilhar fora do sistema.','Se editar após já ter arquivado, use “Atualizar no Drive” para substituir a versão anterior na mesma pasta.'])
h(doc,'Documentos disponíveis',2)
t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
for c,w in zip(t.rows[0].cells,[2.0,4.5]): c.width=Inches(w); shade(c,GREEN); set_cell_margins(c)
for c,txt in zip(t.rows[0].cells,['Documento','Uso recomendado']): font(c.paragraphs[0].add_run(txt),9.5,True,WHITE)
for a,b in [('Ata','Registro formal do encontro, sem a foto.'),('Resumo executivo','Leitura rápida de contexto, temas e resultados.'),('Matriz de ações','O que será feito, por quem, até quando e com qual prioridade.'),('Decisões e bloqueios','Deliberações, pendências e impedimentos.'),('Mapa mental','Estrutura visual automática dos temas e relações, em uma página A4 com fundo branco.')]:
    cells=t.add_row().cells
    for c,w in zip(cells,[2.0,4.5]): c.width=Inches(w); set_cell_margins(c)
    font(cells[0].paragraphs[0].add_run(a),9.2,True,GREEN); font(cells[1].paragraphs[0].add_run(b),9.2,False,INK)
add_text(doc,'A busca superior localiza reuniões, decisões e tarefas. O histórico fica associado à conta e aos dados disponíveis no sistema.',size=9.5)

page(doc); h(doc,'6. Ações e decisões',1)
h(doc,'Ações',2); add_text(doc,'A descrição da ação vem da reunião. Depois do processamento, você pode ajustar os campos operacionais sem alterar o compromisso registrado. Confirme na transcrição se “posso fazer” virou uma responsabilidade efetivamente acordada antes de tratar uma candidatura como compromisso.')
bullets(doc,['Marque a caixa quando a ação estiver concluída.','Clique/edite o responsável.','Defina ou altere o prazo.','Classifique a prioridade.','Depois das alterações, atualize Drive/Trello para refletir a versão vigente.'])
picture(doc,'14-acoes-com-dados.png','Ações reais extraídas de reuniões. Responsável, prazo e prioridade são editáveis.')
h(doc,'Decisões',2); bullets(doc,['Use Decisões para acompanhar deliberações, pendências e bloqueios.','Confirme se o registro corresponde ao que foi efetivamente acordado.','Consulte a reunião de origem antes de compartilhar uma decisão sensível.'])
picture(doc,'06-decisoes.png','Registro consolidado de decisões, pendências e bloqueios.')

page(doc); h(doc,'7. Pergunte à IA',1)
bullets(doc,['Abra “Pergunte à IA”.','Selecione a reunião.','Escreva uma pergunta objetiva e envie.','Leia a resposta semântica, construída com transcrição, resumo, temas, ações e decisões.','Continue perguntando: o histórico mostra cada pergunta e a respectiva resposta.','Para recomeçar, clique em “Apagar conversa” e confirme; somente o histórico da reunião selecionada será removido.'],numbered=True)
h(doc,'Boas perguntas',2); bullets(doc,['Quem participou e quais responsabilidades foram atribuídas?','Quais decisões foram tomadas?','O que deve ser entregue nesta semana?','Quais bloqueios permanecem sem solução?','Resuma a discussão sobre determinado tema.'])
picture(doc,'12-chat-com-dados.png','Chat da reunião: selecione a reunião, pergunte e acompanhe todo o histórico de perguntas e respostas.')
callout(doc,'Escopo','As respostas ficam limitadas ao conteúdo registrado naquela reunião. Confirme informações críticas na transcrição e nos documentos.')
h(doc,'Sair do sistema',2); add_text(doc,'Clique no ícone do KeyNotesAI no topo do menu esquerdo. A sessão será encerrada pela autenticação oficial e o aplicativo retornará à capa. No celular, use o mesmo ícone no cabeçalho.')

page(doc); h(doc,'8. Google Drive',1)
h(doc,'Arquivar ou atualizar',2); bullets(doc,['Conclua a transcrição, a análise e a revisão dos documentos.','Clique em “Arquivar no Drive” na primeira vez.','Nas próximas alterações, clique em “Atualizar no Drive”.','O sistema reutiliza a mesma pasta da reunião e atualiza os arquivos; não deve criar pastas duplicadas.','Após sucesso no Drive, a sincronização do Trello é iniciada automaticamente.'],numbered=True)
callout(doc,'Destino institucional','Conta inovalab.cte@gmail.com → Meu Drive → 4. Arquivos → Reuniões - KeyNotesAI.')
h(doc,'Padrão da pasta',2); add_text(doc,'aaaa-mm-dd : hh:mm - Título da Reunião',bold=True,color=GREEN,size=12)
bullets(doc,['Áudio/gravação, quando incluído.','Ata, resumo executivo, plano de ação, decisões e transcrição em PDF — sem conversão para Google Docs.','Mapa mental em formato vetorial, com fundo branco.','Foto da reunião, se cadastrada.','Links utilizáveis no card do Trello.'])
callout(doc,'Se algo for editado','Salve no KeyNotesAI e execute “Atualizar no Drive”. A pasta existente é a fonte vigente; evite renomeá-la manualmente.')

page(doc); h(doc,'9. Trello',1)
h(doc,'Configuração administrativa',2); add_text(doc,'O destino atual é o quadro INOVALAB - “onde as ideias viram realidade”, lista Reuniões. O status “Conectado” e o destino aparecem em Integrações.')
picture(doc,'15-trello-destino.png','Destino configurado no KeyNotesAI.')
h(doc,'Quando o card é enviado?',2); bullets(doc,['Automaticamente: logo após arquivar ou atualizar a reunião no Google Drive com sucesso.','Manualmente: pelo botão “Sincronizar/Atualizar no Trello”, útil depois de ajustes que ainda não passaram pelo Drive.'])
h(doc,'Abrir o quadro',2); add_text(doc,'Na Visão geral, clique em “Abrir quadro no Trello”. O botão utiliza a URL real do quadro escolhido em Integrações e abre o destino em uma nova aba. Se ainda não houver destino, o sistema orienta o administrador a selecionar Quadro e Lista.')
h(doc,'O que vai no card',2); bullets(doc,['Título e identificação da reunião.','Resumo, decisões, ações, responsáveis, prazos e prioridades.','Links da pasta e dos arquivos no Google Drive.','Checklist das ações.'])
callout(doc,'Sem duplicatas','Cada reunião usa um único card. Novas sincronizações atualizam o mesmo card, desde que o vínculo anterior seja preservado.')

page(doc); h(doc,'10. Administração: usuários e limites',1)
add_text(doc,'Disponível apenas para administradores. O administrador principal é rogerio.bittencourt@ifsc.edu.br.')
h(doc,'Cadastrar usuário',2); bullets(doc,['Abra Administração.','Informe nome e o e-mail exato vinculado à conta ChatGPT do usuário.','Defina o limite de operações por mês.','Clique em “Cadastrar usuário”.','Peça ao usuário que entre com essa mesma conta.'],numbered=True)
h(doc,'Alterar ou remover',2); bullets(doc,['Localize o usuário na lista.','Altere limite e/ou status e salve.','Para revogar o acesso, desative ou elimine o usuário conforme a opção disponível.','Considere o impacto antes de excluir: reuniões do usuário podem continuar armazenadas em integrações institucionais.'])
picture(doc,'08-administracao.png','Administração de usuários, consumo mensal, limites e status.')
callout(doc,'Contagem','Cada operação corresponde a uma transcrição pela OpenAI ou à geração de documentos. O modo híbrido reduz o uso da API para transcrição, mas a análise e a geração continuam contando.')

page(doc); h(doc,'11 — Excluir uma reunião',1)
add_text(doc,'A exclusão completa é uma função administrativa e deve ser usada com cuidado.')
bullets(doc,['Abra Arquivos e selecione a reunião.','Acione a opção administrativa de excluir.','Leia a confirmação, verificando título/data.','Confirme somente se pretende remover o registro local e os artefatos vinculados.','O KeyNotesAI envia a pasta correspondente do Google Drive para a lixeira.','Confirme também o Trello: se o card ainda existir, remova ou arquive-o manualmente.'],numbered=True)
callout(doc,'Recuperação','Itens enviados à lixeira do Drive podem ser recuperados enquanto permanecerem lá. A exclusão dos dados do sistema pode não ser reversível.')
h(doc,'Antes de excluir',2); bullets(doc,['Baixe ou compartilhe o que precisa ser preservado.','Verifique se a reunião correta está selecionada.','Consulte o responsável pela guarda documental, quando aplicável.'])

page(doc); h(doc,'12 — Solução rápida de problemas',1)
t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
for c,w in zip(t.rows[0].cells,[2.25,4.25]): c.width=Inches(w); shade(c,GREEN); set_cell_margins(c)
for c,txt in zip(t.rows[0].cells,['Sintoma','O que fazer']): font(c.paragraphs[0].add_run(txt),9.5,True,WHITE)
problems=[
('Não consigo entrar','Confirme se o e-mail do ChatGPT é exatamente o cadastrado pelo administrador.'),
('Microfone não grava','Revise a permissão do navegador, selecione o microfone correto e recarregue a página.'),
('Meet sem áudio remoto','Use Chrome/Edge, selecione a aba do Meet e marque “Compartilhar áudio da guia”. Não escolha tela inteira ou janela.'),
('Minha voz não aparece','Use “Google Meet + microfone” e autorize o microfone. O modo somente aba pode não incluir sua voz local.'),
('Gravação encerrou sozinha','Interromper o compartilhamento da aba encerra e salva a gravação automaticamente.'),
('“Tirar foto” abre arquivos','Use navegador/dispositivo com câmera disponível e permissão concedida; em desktop sem câmera compatível, use Selecionar arquivo.'),
('Transcrição fraca','Revise áudio, ruído e distância do microfone; corrija a transcrição antes da análise.'),
('Locutor sem nome','Confirme que a pessoa falou o nome na chamada inicial; depois associe manualmente o Locutor A/B na área de conferência.'),
('Vozes trocadas','Corrija a associação na revisão de locutores antes de analisar e gerar documentos.'),
('Sem ações/decisões','Confirme que a reunião foi transcrita e analisada; compromissos precisam estar claros no texto.'),
('Drive não atualiza','Verifique o status em Integrações; tente novamente e confirme a pasta existente.'),
('Trello não atualiza','Primeiro atualize o Drive; depois use o botão manual e confirme quadro/lista conectados.'),
('Limite atingido','Solicite ao administrador a revisão do limite mensal.'),
('Resposta do chat vazia','Selecione uma reunião processada e faça pergunta relacionada ao conteúdo dela.')]
for a,b in problems:
    cells=t.add_row().cells
    for c,w in zip(cells,[2.25,4.25]): c.width=Inches(w); set_cell_margins(c)
    font(cells[0].paragraphs[0].add_run(a),8.8,True,GREEN); font(cells[1].paragraphs[0].add_run(b),8.8,False,INK)

page(doc); h(doc,'Checklist de uso 100%',1)
bullets(doc,['Usuário autorizado e login validado.','Título claro, origem do áudio e modo de processamento escolhidos.','No Google Meet: aba correta selecionada, áudio da guia compartilhado e microfone autorizado quando necessário.','Chamada inicial realizada com uma pessoa dizendo o nome por vez, quando usar identificação de locutores.','Áudio gravado/importado e participantes conferidos.','Associação entre vozes e nomes revisada.','Foto registrada com consentimento, se desejada.','Transcrição revisada.','Análise semântica concluída.','Ata e demais documentos revisados/salvos.','Responsáveis, prazos, prioridades e status das ações ajustados.','Perguntas ao chat verificadas e histórico apagado somente quando desejado.','Pasta do Drive arquivada/atualizada sem duplicidade e com PDFs.','Card do Trello criado/atualizado e links conferidos.','Reunião excluída somente quando realmente necessário.'])
callout(doc,'Fluxo ideal','Google Meet/presencial → gravação → transcrição → revisão → análise → documentos → ações → Drive → Trello.')
h(doc,'Perfis',2)
t=doc.add_table(rows=1,cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
for c,w in zip(t.rows[0].cells,[1.3,2.5,2.7]): c.width=Inches(w); shade(c,GREEN); set_cell_margins(c)
for c,txt in zip(t.rows[0].cells,['Perfil','Pode fazer','Responsabilidade']): font(c.paragraphs[0].add_run(txt),9,True,WHITE)
for row in [('Usuário','Criar/processar reuniões; editar documentos e ações; usar chat e integrações.','Respeitar consentimento, revisar dados e acompanhar seu limite.'),('Administrador','Tudo do usuário; cadastrar/remover pessoas, alterar limites e excluir reuniões.','Governança de acesso, custos e guarda documental.')]:
    cells=t.add_row().cells
    for c,w in zip(cells,[1.3,2.5,2.7]): c.width=Inches(w); set_cell_margins(c)
    for c,txt in zip(cells,row): font(c.paragraphs[0].add_run(txt),8.8,c is cells[0],GREEN if c is cells[0] else INK)
add_text(doc,'Manual versão 1.2, preparado a partir da versão publicada do KeyNotesAI e de telas reais do sistema. A interface poderá receber ajustes sem alterar o fluxo essencial.',color=MUTED,size=8,after=0,align=WD_ALIGN_PARAGRAPH.CENTER)

OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)
